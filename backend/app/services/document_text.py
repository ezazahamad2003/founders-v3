"""Utilities for downloading uploaded files and extracting readable text."""

from __future__ import annotations

import asyncio
import logging
import re
from io import BytesIO
from typing import List, Optional, Sequence
from urllib.parse import quote

import httpx
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from pypdf import PdfReader

from app.config import Settings
from app.models import FileMeta

MAX_CHARS_PER_FILE = 50000  # Increased for legal documents (~8,000 words)
MAX_FILE_SIZE_BYTES = 50 * 1024 * 1024  # 50 MB

logger = logging.getLogger(__name__)


async def build_documents_contexts(
    files_meta: Sequence[FileMeta],
    settings: Settings,
) -> List[str]:
    """Download each file, extract text, and return context snippets for prompts."""
    logger.info(f"[DOC_EXTRACT] build_documents_contexts called with {len(files_meta)} files")
    if not files_meta:
        logger.info(f"[DOC_EXTRACT] No files to process, returning empty")
        return []

    for fm in files_meta:
        logger.info(f"[DOC_EXTRACT] Will process: {fm.original_name} (path={fm.supabase_path}, mime={fm.mime_type})")

    tasks = [_fetch_file_text(file_meta, settings) for file_meta in files_meta]
    results = await asyncio.gather(*tasks, return_exceptions=True)

    contexts: List[str] = []
    for file_meta, result in zip(files_meta, results):
        if isinstance(result, Exception):
            logger.error(f"[DOC_EXTRACT] Failed to read {file_meta.supabase_path}: {result}")
            continue
        snippet = _truncate(result)
        if not snippet:
            logger.warning(f"[DOC_EXTRACT] Empty snippet for {file_meta.original_name}")
            continue
        title = file_meta.original_name or str(file_meta.id)
        logger.info(f"[DOC_EXTRACT] Extracted {len(snippet)} chars from {title}")
        contexts.append(f"### {title}\n{snippet}")
    
    logger.info(f"[DOC_EXTRACT] Returning {len(contexts)} context snippets")
    return contexts


async def _fetch_file_text(file_meta: FileMeta, settings: Settings) -> str:
    """Download bytes from Supabase storage and convert to text."""
    data = await _download_file_from_storage(file_meta, settings)
    return extract_text_from_bytes(
        data=data,
        mime_type=file_meta.mime_type,
        file_name=file_meta.original_name or file_meta.supabase_path,
    )


async def _download_file_from_storage(file_meta: FileMeta, settings: Settings) -> bytes:
    """Download a file from Supabase Storage."""
    if not settings.supabase_project_url:
        logger.error("SUPABASE_PROJECT_URL is not configured.")
        raise RuntimeError("Document storage is not configured.")
    
    # Use service role key for backend operations (bypasses RLS)
    auth_key = settings.supabase_service_role_key_clean or settings.supabase_anon_key_clean
    if not auth_key:
        logger.error("SUPABASE_SERVICE_ROLE_KEY or SUPABASE_ANON_KEY is required to download files.")
        raise RuntimeError("Document storage authentication is missing.")

    base = settings.supabase_project_url.rstrip("/")
    bucket = settings.supabase_storage_bucket_name.strip("/")
    path = file_meta.supabase_path.lstrip("/")
    encoded_path = quote(path, safe="/")
    url = f"{base}/storage/v1/object/{bucket}/{encoded_path}"

    try:
        async with httpx.AsyncClient(timeout=20) as client:
            response = await client.get(
                url,
                headers={
                    "apikey": auth_key,
                    "Authorization": f"Bearer {auth_key}",
                },
            )
            if response.status_code >= 400:
                logger.error(
                    f"Supabase download failed for {file_meta.supabase_path} ({response.status_code}): {response.text}"
                )
                raise RuntimeError(
                    f"Failed to download file from storage (status {response.status_code})"
                )
            
            content_length = len(response.content)
            if content_length > MAX_FILE_SIZE_BYTES:
                logger.warning(
                    f"File {file_meta.supabase_path} exceeds size limit ({content_length} bytes), skipping extraction."
                )
                raise RuntimeError(f"File too large ({content_length / (1024*1024):.1f} MB)")
            
            return response.content
    except httpx.RequestError as exc:
        logger.error(f"Network error downloading {file_meta.supabase_path}: {exc}")
        raise RuntimeError("Network error accessing document storage") from exc


def extract_text_from_bytes(
    *,
    data: bytes,
    mime_type: Optional[str],
    file_name: Optional[str],
) -> str:
    """Convert uploaded file bytes into readable text."""
    if not data:
        return ""

    lower_name = (file_name or "").lower()
    mime = (mime_type or "").lower()

    if _is_plain_text(mime, lower_name):
        return data.decode("utf-8", errors="ignore")

    if _is_pdf(mime, lower_name):
        return _extract_pdf_text(data)

    if _is_docx(mime, lower_name):
        return _extract_docx_text(data)

    if _is_doc(mime, lower_name):
        logger.warning(f"Legacy .doc file not supported: {file_name}")
        return ""

    return ""


def _extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    buffer: List[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text:
            buffer.append(text)
    return "\n".join(buffer)


def _extract_docx_text(data: bytes) -> str:
    try:
        document = DocxDocument(BytesIO(data))
    except (PackageNotFoundError, ValueError, KeyError) as exc:
        logger.error(f"Failed to parse DOCX: {exc}")
        return ""

    parts: List[str] = []
    for paragraph in document.paragraphs:
        content = paragraph.text.strip()
        if content:
            parts.append(content)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _is_plain_text(mime: str, file_name: str) -> bool:
    return mime.startswith("text/") or file_name.endswith(".txt")


def _is_pdf(mime: str, file_name: str) -> bool:
    return mime == "application/pdf" or file_name.endswith(".pdf")


def _is_docx(mime: str, file_name: str) -> bool:
    return (
        mime
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or file_name.endswith(".docx")
    )


def _is_doc(mime: str, file_name: str) -> bool:
    # Legacy .doc files (application/msword) are not supported; only .docx
    return mime == "application/msword" and file_name.endswith(".doc")


def _truncate(text: str) -> str:
    normalized = re.sub(r"\s+", " ", text).strip()
    if not normalized:
        return ""
    if len(normalized) <= MAX_CHARS_PER_FILE:
        return normalized
    return normalized[:MAX_CHARS_PER_FILE] + "…"


