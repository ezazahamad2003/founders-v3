from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from uuid import uuid4

import pytest
from docx import Document

from app.config import Settings
from app.models import FileMeta
from app.services import document_text


def _load_pdf_bytes() -> bytes:
    pdf_path = (
        Path(__file__)
        .resolve()
        .parent.parent
        / "test"
        / "Dawood CV.pdf"
    )
    return pdf_path.read_bytes()


def _build_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("Document Intake Summary", level=1)
    doc.add_paragraph("Client: Founders LLM")
    doc.add_paragraph("Need: Parse DOCX contents for chat grounding.")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_extract_text_from_pdf_contains_expected_sections():
    pdf_bytes = _load_pdf_bytes()
    extracted = document_text.extract_text_from_bytes(
        data=pdf_bytes,
        mime_type="application/pdf",
        file_name="Dawood CV.pdf",
    )
    assert "Muhammad Dawood Saeed" in extracted
    assert "Vision Tech 360" in extracted
    assert "Full Stack Web Developer" in extracted


def test_extract_text_from_docx_contains_expected_sections():
    docx_bytes = _build_docx_bytes()
    extracted = document_text.extract_text_from_bytes(
        data=docx_bytes,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        file_name="intake.docx",
    )
    assert "Document Intake Summary" in extracted
    assert "Client: Founders LLM" in extracted
    assert "Need: Parse DOCX contents" in extracted


@pytest.mark.asyncio
async def test_build_documents_contexts_uses_downloaded_bytes(monkeypatch):
    pdf_bytes = _load_pdf_bytes()

    async def fake_download(file_meta, settings):
        return pdf_bytes

    monkeypatch.setattr(document_text, "_download_file_from_storage", fake_download)

    dummy_file = FileMeta(
        id=uuid4(),
        conversation_id=uuid4(),
        supabase_path="profiles/dawood/Dawood CV.pdf",
        mime_type="application/pdf",
        original_name="Dawood CV.pdf",
        created_at=datetime.now(timezone.utc),
    )

    settings = Settings.model_validate(
        {
            "openai_api_key": "test",
            "supabase_db_url": "postgresql://postgres:postgres@localhost:5432/postgres",
            "supabase_storage_bucket_name": "uploads",
        }
    )

    contexts = await document_text.build_documents_contexts([dummy_file], settings=settings)
    assert len(contexts) == 1
    assert "Muhammad Dawood Saeed" in contexts[0]
    assert contexts[0].startswith("### Dawood CV.pdf")


@pytest.mark.asyncio
async def test_build_documents_contexts_handles_docx(monkeypatch):
    docx_bytes = _build_docx_bytes()

    async def fake_download(file_meta, settings):
        return docx_bytes

    monkeypatch.setattr(document_text, "_download_file_from_storage", fake_download)

    dummy_file = FileMeta(
        id=uuid4(),
        conversation_id=uuid4(),
        supabase_path="profiles/user/intake.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        original_name="intake.docx",
        created_at=datetime.now(timezone.utc),
    )

    settings = Settings.model_validate(
        {
            "openai_api_key": "test",
            "supabase_db_url": "postgresql://postgres:postgres@localhost:5432/postgres",
            "supabase_storage_bucket_name": "uploads",
        }
    )

    contexts = await document_text.build_documents_contexts([dummy_file], settings=settings)
    assert len(contexts) == 1
    assert "Document Intake Summary" in contexts[0]
    assert contexts[0].startswith("### intake.docx")

