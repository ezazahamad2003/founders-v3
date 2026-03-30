"""
document_builder.py
Parses the Option Agreement template and rebuilds a fixed .docx by applying
AI-generated replacements, while preserving every custom Word style.
"""

from __future__ import annotations

import io
from docx import Document
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Template parsing
# ---------------------------------------------------------------------------

def get_template_structure(template_path: str) -> dict:
    """
    Return the ordered list of paragraphs and tables from the template so the
    AI agent can understand what to replace.
    """
    doc = Document(template_path)

    paragraphs = [
        {
            "id": i,
            "style": p.style.name,
            "text": p.text,
            "has_content": bool(p.text.strip()),
        }
        for i, p in enumerate(doc.paragraphs)
    ]

    tables = [
        {
            "id": i,
            "rows": [[cell.text for cell in row.cells] for row in table.rows],
        }
        for i, table in enumerate(doc.tables)
    ]

    return {"paragraphs": paragraphs, "tables": tables}


# ---------------------------------------------------------------------------
# Low-level docx helpers
# ---------------------------------------------------------------------------

_RUN_LIKE_TAGS = frozenset({"r", "hyperlink", "ins", "del"})


def _clear_paragraph(para) -> None:
    """Remove all run-content from a paragraph, keeping the paragraph node."""
    p_elem = para._element
    for child in list(p_elem):
        local = child.tag.split("}")[-1]
        if local in _RUN_LIKE_TAGS:
            p_elem.remove(child)


def _set_paragraph_text(para, text: str) -> None:
    """Replace the paragraph's text with *text*, preserving its Word style."""
    _clear_paragraph(para)
    if text:
        para.add_run(text)


def _set_cell_text(cell, text: str) -> None:
    """Replace the text of a table cell's first paragraph."""
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return
    first_para = cell.paragraphs[0]
    _clear_paragraph(first_para)
    if text:
        first_para.add_run(text)


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def build_document(template_path: str, replacements: dict) -> bytes:
    """
    Open the template (for styles), apply AI-generated replacements, and return
    the finished document as bytes.

    *replacements* must be shaped like::

        {
          "paragraphs": {"<para_id>": "<new text>", ...},
          "tables":     {"<table_id>": [["r0c0", "r0c1", ...], ["r1c0", ...]], ...}
        }
    """
    doc = Document(template_path)

    para_map: dict[int, str] = {
        int(k): v for k, v in replacements.get("paragraphs", {}).items()
    }
    table_map: dict[int, list[list[str]]] = {
        int(k): v for k, v in replacements.get("tables", {}).items()
    }

    # ---- paragraphs --------------------------------------------------------
    for idx, new_text in para_map.items():
        if 0 <= idx < len(doc.paragraphs):
            _set_paragraph_text(doc.paragraphs[idx], new_text)

    # ---- tables ------------------------------------------------------------
    for idx, new_rows in table_map.items():
        if 0 <= idx < len(doc.tables):
            table = doc.tables[idx]
            for row_idx, row in enumerate(table.rows):
                if row_idx >= len(new_rows):
                    break
                for col_idx, cell in enumerate(row.cells):
                    if col_idx < len(new_rows[row_idx]):
                        _set_cell_text(cell, new_rows[row_idx][col_idx])

    # Final safety pass — fix any stray template references the AI may have missed
    _post_process(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Post-processing safety pass
# ---------------------------------------------------------------------------

# Ordered list of (exact_match, replacement) applied across every paragraph
# and table cell after AI replacements. Only catches patterns the AI missed.
_STRING_FIXES: list[tuple[str, str]] = [
    # Most common miss: termination-for-cause paragraph still says Employment Agreement
    ("termination of the Employment Agreement for Cause", "termination of the Consulting Agreement for Cause"),
    ("the Employment Agreement for Cause", "the Consulting Agreement for Cause"),
    ("Employment Agreement for Cause", "Consulting Agreement for Cause"),
    # Generic Employment Agreement references
    ("the Employment Agreement", "the Consulting Agreement"),
    ("an Employment Agreement", "a Consulting Agreement"),
    ("Employment Agreement", "Consulting Agreement"),
    # employment relationship references
    ("employment with the Corporation", "consulting engagement with the Corporation"),
    ("Optionee\u2019s employment", "Optionee\u2019s consulting engagement"),
    ("Optionee's employment", "Optionee's consulting engagement"),
    # Leftover angle-star placeholders (Resignation & Release template style)
    ("<*>", "[TO BE COMPLETED]"),
]


def _fix_text(text: str) -> str:
    for old, new in _STRING_FIXES:
        text = text.replace(old, new)
    return text


def _post_process(doc) -> None:
    """Apply string-level fixes across all paragraphs and table cells."""
    for para in doc.paragraphs:
        original = para.text
        fixed = _fix_text(original)
        if fixed != original:
            _set_paragraph_text(para, fixed)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original = para.text
                    fixed = _fix_text(original)
                    if fixed != original:
                        _clear_paragraph(para)
                        if fixed:
                            para.add_run(fixed)
