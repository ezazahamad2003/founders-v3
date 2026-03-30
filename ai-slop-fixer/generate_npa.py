from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

def set_font(run, size=11, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold

def heading(text, level=1, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=12 if level == 1 else 11, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(text)
    set_font(run, size=11)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def section_header(num, title):
    p = doc.add_paragraph()
    run = p.add_run(f"{num}.  {title}")
    set_font(run, size=11, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def subsection_header(num, title):
    p = doc.add_paragraph()
    run = p.add_run(f"{num}  {title}")
    set_font(run, size=11, bold=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    return p

def toc_row(num, title, pg=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0 if "." not in num else 0.25)
    run = p.add_run(f"{num}\t{title}")
    set_font(run, size=10)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

# ============================================================
# COVER
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("[COMPANY NAME], INC.")
set_font(r, size=14, bold=True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("NOTE PURCHASE AGREEMENT")
set_font(r2, size=14, bold=True)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Dated as of __________, 2026")
set_font(r3, size=12)

doc.add_paragraph()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
heading("TABLE OF CONTENTS", level=1, center=True)

toc_entries = [
    ("1.", "Definitions"),
    ("2.", "Amount and Terms of the Notes"),
    ("2.1", "Issuance and Repayment of Notes"),
    ("2.2", "Conversion of the Notes"),
    ("2.3", "Payment in Lieu of Conversion"),
    ("3.", "Closing Mechanics"),
    ("3.1", "Closing"),
    ("3.2", "Covenants"),
    ("4.", "Representations and Warranties of the Company"),
    ("4.1", "Organization, Good Standing and Qualification"),
    ("4.2", "Authorization"),
    ("4.3", "Compliance with Other Instruments"),
    ("4.4", "Valid Issuance of Stock"),
    ("4.5", "Litigation"),
    ("4.6", "Disclosure"),
    ("4.7", "Securities Law Compliance"),
    ("4.8", "Use of Proceeds"),
    ("5.", "Representations and Warranties of the Lender"),
    ("5.1", "Authorization"),
    ("5.2", "Purchase Entirely for Own Account"),
    ("5.3", "Disclosure of Information; No General Solicitation"),
    ("5.4", "Investment Experience"),
    ("5.5", "Accredited Investor"),
    ("5.6", "Exculpation Among Lenders"),
    ("5.7", "Restricted Securities"),
    ("5.8", "Further Limitations on Disposition"),
    ("5.9", "Legends"),
    ("5.10", "Confidentiality"),
    ("6.", "Miscellaneous"),
    ("6.1", "Representations and Warranties Survive the Closing; Severability"),
    ("6.2", "Most Favored Nation"),
    ("6.3", "Pro Rata Rights"),
    ("6.4", "Successors and Assigns"),
    ("6.5", "Governing Law; Submission to Jurisdiction"),
    ("6.6", "Counterparts"),
    ("6.7", "Facsimile Signatures"),
    ("6.8", "Titles and Headings"),
    ("6.9", "Notices"),
    ("6.10", "Finder\u2019s Fee"),
    ("6.11", "Entire Agreement; Amendments and Waivers"),
    ("6.12", "Effect of Amendment or Waiver"),
    ("6.13", "Stock Purchase Agreement"),
    ("6.14", "Further Assurances"),
    ("6.15", "Expenses"),
]
for num, title in toc_entries:
    toc_row(num, title)

doc.add_page_break()

# ============================================================
# AGREEMENT BODY
# ============================================================
heading("NOTE PURCHASE AGREEMENT", level=1, center=True)

body(
    'THIS NOTE PURCHASE AGREEMENT (this \u201cAgreement\u201d) is made as of __________, 2026, by and among '
    '[COMPANY NAME], INC., a New York corporation (the \u201cCompany\u201d), and the lender(s) named on the '
    'Signature Page hereof (each a \u201cLender\u201d and collectively, the \u201cLenders\u201d).  '
    'Capitalized terms not otherwise defined in this Agreement shall have the meanings given to them in Section\u00a01 below.'
)

body(
    'WHEREAS, the Lender intends to provide certain Consideration to the Company as described for the Lender on the '
    'Signature Page hereof;'
)
body(
    'WHEREAS, the parties wish to provide for the sale and issuance of Notes in return for the provision by the Lender '
    'of the Consideration to the Company; and'
)
body(
    'WHEREAS, the Note shall be convertible into equity securities of the Company, or the outstanding principal of the '
    'Notes, plus all accrued and unpaid interest on the outstanding principal of the Notes, shall be payable in cash, '
    'in each case, in accordance with the terms set forth herein and in the Note.'
)
body('NOW, THEREFORE, the parties hereby agree as follows:')

# ===== SECTION 1 =====
section_header("1", "Definitions.")

definitions = [
    ('"Board"', 'means the Board of Directors of the Company.'),
    ('"Common Stock"', 'means common stock, par value $0.0001 per share, of the Company.'),
    ('"Consideration"', 'shall mean the amount of money paid by the Lender pursuant to this Agreement as shown on the Signature Page hereof.'),
    ('"Conversion Price"', 'shall mean, as applicable:\n'
     '(a) with respect to a conversion pursuant to Section\u00a02.2(a) below, a price per Conversion Share that reflects a pre-conversion Company valuation equal to the Valuation Cap;\n'
     '(b) with respect to a conversion pursuant to Section\u00a02.2(b)(i) or Section\u00a02.2(b)(iii) below, the lower of (i) the Fair Market Value per Conversion Share, or (ii) a price per Conversion Share based on the Valuation Cap;\n'
     '(c) with respect to a conversion pursuant to Section\u00a02.2(b)(ii) below, the lower of either (i) the product of (x) the Discount Rate multiplied by (y) the price paid per share for Equity Securities by the investors at the last closing that is part of the Next Equity Financing, or (ii) a price per Equity Security based on the Valuation Cap; and\n'
     '(d) with respect to a conversion pursuant to Section\u00a02.2(b)(iv) below, the lower of either (i) the product of (x) the Discount Rate multiplied by (y) the rate proportional to the same valuation for the Conversion Shares accorded thereto in the Corporate Transaction, or (ii) a price per Conversion Share based on the Valuation Cap.'),
    ('"Conversion Shares"', 'shall mean, for purposes of determining the type of securities issuable upon conversion of the Note, (i) if the Note is converted to equity pursuant to Section\u00a02.2(b)(ii) below, the Equity Securities issued in the Next Equity Financing, or (ii) if the Note is converted (or deemed converted) to equity pursuant to Section\u00a02.2(a), Section\u00a02.2(b)(i), Section\u00a02.2(b)(iii) or Section\u00a02.2(b)(iv) below, shares of Common Stock.'),
    ('"Corporate Transaction"', 'shall include (A) the closing of the sale, transfer or other disposition of all or substantially all of the Company\u2019s assets, (B) the consummation of the merger or consolidation of the Company with or into another entity (except a merger or consolidation in which the holders of capital stock of the Company immediately prior to such merger or consolidation continue to hold at least a majority of the voting power of the capital stock of the Company or the surviving or acquiring entity), (C) the closing of the transfer (whether by merger, consolidation or otherwise), in one transaction or a series of related transactions, to a person or group of affiliated persons (other than an underwriter of the Company\u2019s securities), of the Company\u2019s securities if, after such closing, such person or group of affiliated persons would hold a majority or more of the outstanding voting power of the capital stock of the Company (or the surviving or acquiring entity), or (D) a liquidation, dissolution or winding up of the Company; provided, however, that a transaction shall not constitute a Corporate Transaction if its sole purpose is to change the state of the Company\u2019s incorporation or to create a holding company that will be owned in substantially the same proportions by the persons who held the Company\u2019s securities immediately prior to such transaction.'),
    ('"Discount Rate"', 'shall mean [__]% ([.__]), representing the percentage of the price per Equity Security in the Next Equity Financing at which the Note shall convert, such that the effective discount to the Lender is equal to one minus the Discount Rate.'),
    ('"Equity Securities"', 'shall mean, whether Common Stock or a newly authorized class or series of preferred stock of the Company, the securities issued by the Company to purchasers in the Next Equity Financing.'),
    ('"Event of Default"', 'is as defined under the terms of the Note.'),
    ('"Fair Market Value"', 'shall mean the fair market value of the relevant securities, as determined in good faith by the Board; provided, that in the case of a Qualified IPO, the fair market value of a share of Common Stock is the price at which shares of the Common Stock are sold to the public in such Qualified IPO.'),
    ('"Lenders"', 'shall mean the persons purchasing Notes under the terms of the Note Purchase Agreements.'),
    ('"Majority Note Holders"', 'shall mean the holders of a majority in interest of the aggregate outstanding principal amount of Notes.'),
    ('"Maturity Date"', 'shall mean, with respect to the Notes, [________] (being approximately [18\u201324] months from the Issue Date).'),
    ('"Next Equity Financing"', 'shall mean the next sale (or series of related sales) by the Company of its Equity Securities following the date of the last issuance of any of the Notes, from which sale of Equity Securities the Company receives gross proceeds of not less than $[___] (not including the conversion of the Notes, but including the principal amount of any other indebtedness that is converted).'),
    ('"Note"', 'shall mean one or more promissory notes issued to the Lender pursuant to Section\u00a02.1 below, the form of which is attached hereto as Schedule\u00a0A.'),
    ('"Notes"', 'shall mean all of the convertible promissory notes issued as part of a series of such convertible promissory notes to the Lenders under the Note Purchase Agreements, each of which shall be in the form of the Note.'),
    ('"Note Purchase Agreements"', 'shall mean this Agreement and any other equivalent agreements with respect to the purchase and issuance of the Notes as part of the Offering.'),
    ('"Offering"', 'shall mean the sale and issuance of the Notes under the following terms: each Note to bear terms equivalent to the terms of the Note; all of the Notes together to equal up to a maximum of $[___] in aggregate principal, unless such aggregate principal amount is increased by the Company in its discretion; and all of the Notes to be sold pursuant to the terms of the Note Purchase Agreements.'),
    ('"Qualified IPO"', 'means the sale of shares of Common Stock in an underwritten public offering pursuant to an effective registration statement under the Securities Act.'),
    ('"Securities Act"', 'means the Securities Act of 1933, as amended.'),
    ('"Valuation Cap"', 'shall mean a pre-money valuation of $[___], which shall be the maximum Company valuation used to calculate the Conversion Price.'),
]

for term, defn in definitions:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_term = p.add_run(term + "  ")
    set_font(r_term, bold=True)
    r_def = p.add_run(defn)
    set_font(r_def)

# ===== SECTION 2 =====
section_header("2", "Amount and Terms of the Notes.")

subsection_header("2.1", "Issuance and Repayment of Notes.")
body(
    '(a)\tIssuance.  In return for the Consideration paid by the Lender, the Company shall sell and issue to such Lender '
    'one or more Notes.  Each Note shall have an initial principal balance equal to that portion of the Consideration paid by '
    'such Lender for the Note(s), as set forth on the Signature Page hereof.  Each Note shall be convertible into Conversion '
    'Shares pursuant to Section\u00a02.2 below.  Each Note is part of a series of Notes sold pursuant to the Note Purchase '
    'Agreements, all of which are pari passu in right of payment.'
)
body(
    '(b)\tRepayment.  To the extent that the Note remains outstanding and shall not previously have been converted into '
    'Conversion Shares in accordance with Section\u00a02.2 below, to the extent that there exists an outstanding Event of '
    'Default, absent an election to convert the Note by either the Note holder or the Majority Note Holders, the Company '
    'shall pay the applicable Note holders an amount in cash pursuant to Section\u00a02.3 below.'
)

subsection_header("2.2", "Conversion of the Notes.")

body('(a)\tOptional Conversion.')
body(
    'At any time upon the occurrence and continuation of an Event of Default that has not been cured or waived, at the '
    'election of either the Note holder (for itself) or the Majority Note Holders for all Lenders collectively, outstanding '
    'principal of the Note and all accrued and unpaid interest thereon will be converted into the applicable number of '
    'Conversion Shares, at the applicable Conversion Price.', indent=True
)

body('(b)\tMandatory Conversion.')
body(
    '(i)\tMaturity.  To the extent that there is not then an outstanding Event of Default under the terms of the Notes, on '
    'the Maturity Date, the outstanding principal of the Note and all accrued and unpaid interest thereon will be automatically '
    'converted pursuant to this Section\u00a02.2(b)(i) at the applicable Conversion Price into the applicable number of '
    'Conversion Shares, to the extent not previously converted.  Not less than the aggregate of (x) the outstanding principal '
    'of all such Note holder\u2019s Notes and (y) all accrued and unpaid interest thereon will be automatically converted into '
    'Conversion Shares, as such principal and interest are determined as of the Maturity Date.  The number of Conversion Shares '
    'issuable upon such conversion shall be equal to the quotient of (y) the aggregate of the outstanding principal of such '
    'Note and all accrued and unpaid interest thereon through the date of conversion, divided by (z) the applicable Conversion '
    'Price.  Notwithstanding the foregoing, if there is an outstanding Event of Default on the Maturity Date, the holder may '
    'instead elect to receive repayment of principal plus accrued interest in cash.', indent=True
)
body(
    '(ii)\tNext Equity Financing.  In connection with the Next Equity Financing, the Note and all accrued and unpaid interest '
    'thereon will be automatically converted into Conversion Shares upon the closing of the Next Equity Financing.  The number '
    'of Conversion Shares to be issued upon such conversion shall be equal to the quotient of (x) the aggregate of the '
    'outstanding principal of the Note, plus all accrued and unpaid interest thereon through the date of conversion, divided '
    'by (y) the applicable Conversion Price.  At least seven (7) days prior to the closing of the Next Equity Financing, the '
    'Company shall notify the holder of the Note in writing of the terms under which the Equity Securities of the Company will '
    'be sold in such Next Equity Financing.  The issuance of Conversion Shares pursuant to the conversion of the Note shall be '
    'upon and subject to the same terms and conditions applicable to the Equity Securities sold in the Next Equity Financing.', indent=True
)
body(
    '(iii)\tQualified IPO.  The outstanding principal of the Note and all accrued and unpaid interest thereon will be '
    'automatically converted into Conversion Shares upon the closing of the Qualified IPO.  The number of Conversion Shares '
    'to be issued upon such conversion shall be equal to the quotient of (x) the aggregate of the outstanding principal of '
    'the Note, plus all accrued and unpaid interest thereon through the date of conversion, divided by (y) the applicable '
    'Conversion Price.  At least seven (7) days prior to the closing of the Qualified IPO, the Company shall notify the '
    'holder of the Note in writing of the terms of the Qualified IPO.', indent=True
)
body(
    '(iv)\tCorporate Transaction.  All outstanding principal and accrued and unpaid interest thereon will be automatically '
    'converted into Conversion Shares upon the closing of a Corporate Transaction.  The number of Conversion Shares to be '
    'issued upon such conversion shall be equal to the quotient of (x) the aggregate of the outstanding principal of the '
    'Note, plus all accrued and unpaid interest thereon through the date of conversion, divided by (y) the applicable '
    'Conversion Price.  At least seven (7) days prior to the closing of the Corporate Transaction, the Company shall notify '
    'the holder of the Note in writing of the terms under which the Common Stock will be sold in such Corporate Transaction.  '
    'Any conversion of the Note pursuant to this Section\u00a02.2(b)(iv) shall be deemed to have been made immediately prior '
    'to the Corporate Transaction, and on and after such date the persons entitled to receive the securities issuable upon such '
    'conversion shall be treated for all purposes as the record holders of such securities.  Notwithstanding the foregoing, '
    'upon the occurrence of a Corporate Transaction prior to conversion, the holder may instead elect to receive repayment '
    'of principal plus accrued interest in cash.', indent=True
)

body(
    '(c)\tNo Fractional Shares.  Upon the conversion of the Note into Conversion Shares, in lieu of any fractional shares to '
    'which the holder of the Note would otherwise be entitled, the amount of Conversion Shares issuable to the holder of the '
    'Note shall be rounded to the nearest whole number.'
)
body(
    '(d)\tMechanics and Effect of Conversion.  The Company shall not be required to issue or deliver the Conversion Shares '
    'until the holder of the Note has surrendered the Note to the Company in accordance with the terms of this '
    'Section\u00a02.2(d).  Any conversion pursuant to Section\u00a02.2 may be made contingent upon the closing of the Next '
    'Equity Financing, the relevant Corporate Transaction or the Qualified IPO, as applicable.  Upon conversion of the Note '
    'pursuant to the terms of this Agreement, the Company will be forever released from all of its obligations and liabilities '
    'under such Note, including, without limitation, with respect to any obligations regarding payment of any principal and '
    'accrued interest.  Upon conversion of a Note pursuant to Section\u00a02.2(b)(ii) or Section\u00a02.2(b)(iv), as '
    'applicable, the Note holder agrees to execute and deliver such transaction documents entered into by other purchasers '
    'participating in the Next Equity Financing or Corporate Transaction as shall be determined by the Board of Directors of '
    'the Company, acting reasonably (which may include a purchase agreement, an investor rights agreement and other ancillary '
    'agreements as may be required), with customary representations and warranties and transfer restrictions (including, '
    'without limitation, a 180-day lock-up agreement in connection with an initial public offering).  The Note holder also '
    'agrees to deliver the original of such Note for cancellation at the earlier of the closing of the Next Equity Financing, '
    'Corporate Transaction, or Qualified IPO, as applicable, or within five (5) days of the Company giving written notice of '
    'conversion; provided, however, that upon such closing, the Note shall be deemed converted and of no further force and '
    'effect, whether or not delivered for cancellation.  The Company shall, as soon as practicable thereafter, issue and '
    'deliver to the Note holder a certificate or certificates for the number and type of securities to which such Note holder '
    'shall be entitled upon such conversion.'
)
body(
    '(e)\tPayment of Interest in Cash.  Notwithstanding the foregoing, accrued and unpaid interest may be paid in cash at '
    'the option of the Company.'
)

subsection_header("2.3", "Payment in Lieu of Conversion.")
body(
    'In the event that (a) following an Event of Default, the Majority Note Holders accelerate payment of the Note, or '
    '(b) the Note remains outstanding on the Maturity Date, and there is an outstanding Event of Default under the terms '
    'of the Note, then with respect to each Note, unless the Company is given a notice of election to convert the Note in '
    'accordance with Section\u00a02.2(a) and Section\u00a02.2(d) hereunder from the Majority Note Holders for all Notes, '
    'or from a Note holder on its own behalf with respect to its Note, the Company shall immediately pay in cash the amount '
    'of outstanding principal of such Note plus all accrued and unpaid interest on the outstanding principal of such Note, '
    'or on all Notes, as applicable, as determined as of the later of (i) the date of acceleration following such Event of '
    'Default under the terms of the Note, or (ii) if the Note remains unpaid and no election to convert the Note is given '
    'pursuant to this Section\u00a02.3 on or prior to the Maturity Date, the Maturity Date.  The Note holder also agrees to '
    'surrender the original of such Note prior to its receipt of payment by the Company.  Upon payment of a Note by the '
    'Company pursuant to the terms of this Agreement, the Company will be forever released from all of its obligations and '
    'liabilities under such Note, including, without limitation, with respect to any obligations regarding payment of any '
    'principal and accrued interest.'
)

# ===== SECTION 3 =====
section_header("3", "Closing Mechanics.")

subsection_header("3.1", "Closing.")
body(
    'The closing (the \u201cClosing\u201d) of the purchase of the Note in return for the Consideration paid by each Lender '
    'shall take place remotely via the exchange of documents and signatures, at 10:00\u00a0a.m., on __________, 2026, or at '
    'such other time and place as the Company and the Lender agree upon orally or in writing.  At the Closing, the Lender '
    'shall deliver the respective Consideration to the Company and the Company shall deliver to the Lender one or more '
    'executed Notes in return for the respective Consideration provided to the Company.'
)

subsection_header("3.2", "Covenants.")
body(
    'In the event of mandatory or optional conversion of the Note pursuant to any of Section\u00a02.2(a) or '
    'Section\u00a02.2(b) hereunder, the Company shall take such corporate action as shall be reasonably necessary, '
    'including, without limitation, amending its certificate of incorporation to increase its authorized equity interests '
    '(and equity interests for issuance on conversion of any such new securities) to such number and class of equity '
    'interests as shall be sufficient to convert this Note pursuant to any of Section\u00a02.2(a) or Section\u00a02.2(b) '
    'hereof, prior to the effective date of such conversion.'
)

# ===== SECTION 4 =====
section_header("4", "Representations and Warranties of the Company.")
body('In connection with the transactions provided for herein, the Company hereby represents and warrants to the Lenders that:')

subsection_header("4.1", "Organization, Good Standing and Qualification.")
body(
    'The Company is a corporation duly organized, validly existing, and in good standing under the laws of the State of '
    'New York and has all requisite corporate power and authority to carry on its business as now conducted.  The Company '
    'is duly qualified to transact business and is in good standing in each jurisdiction in which the failure to so qualify '
    'would have a material adverse effect on its business or properties.'
)

subsection_header("4.2", "Authorization.")
body(
    'Except for the authorization and issuance of the shares issuable in connection with the Next Equity Financing, a '
    'Qualified IPO, or any Corporate Transaction, all corporate action has been taken on the part of the Company and its '
    'officers, directors and stockholders necessary for the authorization, execution and delivery of this Agreement and the '
    'Notes.  This Agreement constitutes the Company\u2019s valid and legally binding obligation, enforceable in accordance '
    'with its terms, except as may be limited by (i) applicable bankruptcy, insolvency, reorganization, or similar laws '
    'relating to or affecting the enforcement of creditors\u2019 rights and (ii) laws relating to the availability of '
    'specific performance, injunctive relief or other equitable remedies.'
)

subsection_header("4.3", "Compliance with Other Instruments.")
body(
    'Neither the authorization, execution and delivery of this Agreement, nor the issuance and delivery of the Notes, will '
    'constitute or result in a material default or violation of any law or regulation applicable to the Company or any '
    'material term or provision of the Company\u2019s current certificate of incorporation or bylaws or any material '
    'agreement or instrument by which it is bound or to which its properties or assets are subject.'
)

subsection_header("4.4", "Valid Issuance of Stock.")
body(
    'The Conversion Shares to be issued, sold and delivered upon conversion of the Notes will be duly authorized and validly '
    'issued, fully paid and nonassessable and, based in part upon the representations and warranties of the Lenders in this '
    'Agreement, will be issued in compliance with all applicable federal and state securities laws.'
)

subsection_header("4.5", "Litigation.")
body(
    'There is no claim, action, suit, proceeding, arbitration, formal or written complaint, charge or investigation pending '
    'or, to the Company\u2019s knowledge, currently threatened against the Company, which to the knowledge of the Company '
    'after reasonable inquiry, would reasonably be expected to cause a material adverse effect on the business, assets '
    '(including intangible assets), liabilities, financial condition, property, prospects or results of operations of the '
    'Company, or that questions the validity of the Note Purchase Agreements or the Notes or the right of the Company to '
    'enter into any of such agreements, or to consummate the transactions contemplated hereby or thereby.  The Company is '
    'not a party or subject to the provisions of any order, writ, injunction, judgment or decree of any court or government '
    'agency or instrumentality.  There is no action, suit, proceeding or investigation by the Company currently pending or '
    'which the Company intends to initiate.'
)

subsection_header("4.6", "Disclosure.")
body(
    'The Company has made available to the Lenders all the information reasonably available to the Company that the Lenders '
    'have requested for deciding whether to acquire the Notes.  No representation or warranty of the Company contained in '
    'this Agreement contains any untrue statement of a material fact or, to the Company\u2019s knowledge, omits to state a '
    'material fact necessary in order to make the statements contained herein or therein not misleading in light of the '
    'circumstances under which they were made.'
)

subsection_header("4.7", "Securities Law Compliance.")
body(
    'The offer and sale of the Notes is being made pursuant to an exemption from registration under the Securities Act.  '
    'The Company intends to rely on Section\u00a04(a)(2) of the Securities Act and/or Regulation D promulgated thereunder, '
    'as well as applicable New York State securities law exemptions.  Subject to the accuracy of the Lenders\u2019 '
    'representations and warranties in Section\u00a05, the offer, sale and issuance of the Notes do not require '
    'registration under the Securities Act or applicable New York State securities laws.'
)

subsection_header("4.8", "Use of Proceeds.")
body(
    'The Company shall use the proceeds from the sale of the Notes for working capital and general corporate purposes.'
)

# ===== SECTION 5 =====
section_header("5", "Representations and Warranties of the Lender.")
body('In connection with the transactions provided for herein, the Lender hereby represents and warrants to the Company that:')

subsection_header("5.1", "Authorization.")
body(
    'This Agreement constitutes such Lender\u2019s valid and legally binding obligation, enforceable in accordance with its '
    'terms, except as may be limited by (i) applicable bankruptcy, insolvency, reorganization, or similar laws relating to '
    'or affecting the enforcement of creditors\u2019 rights and (ii) laws relating to the availability of specific '
    'performance, injunctive relief or other equitable remedies.  Each Lender represents that it has full power and '
    'authority to enter into this Agreement.'
)

subsection_header("5.2", "Purchase Entirely for Own Account.")
body(
    'The Lender acknowledges that this Agreement is made with such Lender in reliance upon such Lender\u2019s '
    'representation to the Company that the Notes, the Conversion Shares and any Common Stock issuable upon conversion '
    'of the Conversion Shares (collectively, the \u201cSecurities\u201d) will be acquired for investment for such '
    'Lender\u2019s own account, not as a nominee or agent, and not with a view to the resale or distribution of any '
    'part thereof, and that such Lender has no present intention of selling, granting any participation in, or otherwise '
    'distributing the same.  By executing this Agreement, such Lender further represents that such Lender does not have '
    'any contract, undertaking, agreement or arrangement with any person to sell, transfer or grant participations to such '
    'person or to any third person, with respect to the Securities.'
)

subsection_header("5.3", "Disclosure of Information; No General Solicitation.")
body(
    'Such Lender acknowledges that it has received all the information it considers necessary or appropriate for deciding '
    'whether to acquire the Securities.  Such Lender further represents that it has had an opportunity to ask questions '
    'and receive answers from the Company regarding the terms and conditions of the offering of the Securities.'
)
body(
    'The Securities were not offered to the Lender through, and the Lender is not aware of, any form of general '
    'solicitation or general advertising, including, without limitation, (i) any advertisement, article, notice or other '
    'communication published in any newspaper, magazine or similar media or broadcast over television or radio, and '
    '(ii) any seminar or meeting whose attendees have been invited by any general solicitation or general advertising.'
)

subsection_header("5.4", "Investment Experience.")
body(
    'Such Lender is an investor in securities of companies in the development stage and acknowledges that it is able to '
    'fend for itself, can bear the economic risk of its investment and has such knowledge and experience in financial or '
    'business matters that it is capable of evaluating the merits and risks of the investment in the Securities.  If other '
    'than an individual, such Lender also represents that it has not been organized solely for the purpose of acquiring '
    'the Securities.'
)

subsection_header("5.5", "Accredited Investor.")
body(
    'Such Lender is an \u201caccredited investor\u201d within the meaning of Rule\u00a0501 of Regulation\u00a0D '
    'promulgated under the Securities Act, as presently in effect.'
)

subsection_header("5.6", "Exculpation Among Lenders.")
body(
    'Such Lender is not relying upon any person, firm, corporation or stockholder, other than the Company and its '
    'officers and directors in their capacities as such, in making its investment or decision to invest in the Company, '
    'and agrees that no other Lender, nor the respective controlling persons, officers, directors, partners, agents, '
    'stockholders or employees of any other Lender, shall be liable for any action heretofore or hereafter taken or '
    'omitted to be taken by any of them in connection with the purchase and sale of the Notes or the Conversion Shares.'
)

subsection_header("5.7", "Restricted Securities.")
body(
    'Such Lender understands that the Securities have not been, and will not be, registered under the Securities Act, '
    'by reason of a specific exemption from the registration provisions of the Securities Act which depends upon, among '
    'other things, the bona fide nature of the investment intent and the accuracy of such Lender\u2019s representations '
    'as expressed herein.  Such Lender understands that the Securities are \u201crestricted securities\u201d under '
    'applicable United States federal and state securities laws and that, pursuant to these laws, such Lender must hold '
    'the Securities indefinitely unless they are registered with the Securities and Exchange Commission (the \u201cSEC\u201d) '
    'and qualified by state authorities, or an exemption from such registration and qualification requirements is available.'
)

subsection_header("5.8", "Further Limitations on Disposition.")
body(
    'Without in any way limiting the representations and warranties set forth above, such Lender further agrees not to '
    'make any disposition of all or any portion of the Securities unless and until the transferee has agreed in writing '
    'for the benefit of the Company to be bound by this Section\u00a05 and:'
)
body(
    '(a)\tThere is then in effect a registration statement under the Securities Act covering such proposed disposition '
    'and such disposition is made in accordance with such registration statement; or', indent=True
)
body(
    '(b)\tSuch Lender shall have notified the Company of the proposed disposition and shall have furnished the Company '
    'with a detailed statement of the circumstances surrounding the proposed disposition and, if reasonably requested '
    'by the Company, such Lender shall have furnished the Company with an opinion of counsel reasonably satisfactory '
    'to the Company that such disposition will not require registration of such shares under the Securities Act.', indent=True
)
body(
    'In addition, the Lender agrees that it shall not make any disposition of any Note, Conversion Share or other '
    'Securities to any of the Company\u2019s competitors, as determined by the Company in good faith.'
)

subsection_header("5.9", "Legends.")
body(
    'Such Lender understands and agrees that the Company will place the legend set forth below or a similar legend on '
    'any certificate evidencing any Securities, together with any other legends that may be required by state or federal '
    'securities laws, the Company\u2019s certificate of incorporation or bylaws, any other agreement between such Lender '
    'and the Company or any agreement between such Lender and any third party:'
)
p_legend = doc.add_paragraph()
p_legend.paragraph_format.left_indent = Inches(0.5)
p_legend.paragraph_format.right_indent = Inches(0.5)
p_legend.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r_legend = p_legend.add_run(
    '\u201cTHE SECURITIES REPRESENTED HEREBY HAVE BEEN ACQUIRED FOR INVESTMENT AND HAVE NOT BEEN REGISTERED UNDER THE '
    'SECURITIES ACT OF 1933, AS AMENDED (THE \u201cACT\u201d), OR UNDER ANY STATE SECURITIES LAW.  THE SECURITIES MAY '
    'NOT BE OFFERED, SOLD, OR OTHERWISE TRANSFERRED UNLESS THEY ARE REGISTERED UNDER THE ACT AND ALL APPLICABLE STATE '
    'SECURITIES LAWS OR ARE IN COMPLIANCE WITH AN EXEMPTION THEREFROM.  THE SECURITIES REPRESENTED BY THIS CERTIFICATE '
    'ALSO ARE SUBJECT TO AND MAY NOT BE OFFERED, SOLD, TRANSFERRED, ENCUMBERED OR OTHERWISE DISPOSED OF EXCEPT UPON '
    'SATISFACTION OF CERTAIN CONDITIONS SET FORTH IN A CERTAIN NOTE PURCHASE AGREEMENT DATED AS OF __________, 2026, '
    'BETWEEN THE MAKER AND THE PURCHASERS NAMED THEREIN.\u201d'
)
set_font(r_legend, size=10)

subsection_header("5.10", "Confidentiality.")
body(
    'Each Lender agrees that such Lender will keep confidential and will not disclose, divulge, or use for any purpose '
    '(other than to monitor its investment in the Company) any confidential information obtained from the Company in '
    'connection with this Agreement, unless such confidential information (a) is known or becomes known to the public '
    'in general (other than as a result of a breach of this Section\u00a05.10 by such Lender), (b) is or has been '
    'independently developed or conceived by such Lender without use of the Company\u2019s confidential information, '
    'or (c) is or has been made known or disclosed to such Lender by a third party without a breach of any obligation '
    'of confidentiality such third party may have to the Company; provided, however, that a Lender may disclose '
    'confidential information (i) to its attorneys, accountants, consultants, and other professionals to the extent '
    'necessary to obtain their services in connection with monitoring its investment in the Company; (ii) to any '
    'prospective purchaser of a Note from such Lender, if such prospective purchaser agrees to be bound by the '
    'provisions of this Section\u00a05.10; (iii) to any partner, member, stockholder, or wholly owned subsidiary '
    '(each, an \u201cAffiliate\u201d) of such Lender in the ordinary course of business, provided that such Lender '
    'informs such Affiliate that such information is confidential and directs such Affiliate to maintain the '
    'confidentiality of such information; or (iv) as may otherwise be required by law, regulation, rule, court order '
    'or subpoena, provided that such Lender promptly notifies the Company of such disclosure and takes reasonable steps '
    'to minimize the extent of any such required disclosure.'
)

# ===== SECTION 6 =====
section_header("6", "Miscellaneous.")

subsection_header("6.1", "Representations and Warranties Survive the Closing; Severability.")
body(
    'The Company\u2019s and the Lender\u2019s representations and warranties shall survive the Closing '
    'notwithstanding any due diligence investigation made by or on behalf of the party seeking to rely thereon.  '
    'In the event that any provision of this Agreement becomes or is declared by a court of competent jurisdiction '
    'to be illegal, unenforceable or void, this Agreement shall continue in full force and effect without said '
    'provision; provided, that no such severance shall be effective if it materially changes the economic benefit '
    'of this Agreement to any party.'
)

subsection_header("6.2", "Most Favored Nation.")
body(
    'If the Company issues any subsequent convertible promissory note prior to the conversion or repayment of this '
    'Note that contains terms more favorable to the holder thereof than the terms of this Note (including, without '
    'limitation, a lower valuation cap, a higher discount rate, or a lower interest rate), the Lender shall have '
    'the right, exercisable within thirty (30) days of written notice from the Company of such issuance, to amend '
    'this Note to include such more favorable terms.  This Section\u00a06.2 shall not apply to differences solely '
    'in principal amount.'
)

subsection_header("6.3", "Pro Rata Rights.")
body(
    'In the event that prior to the Maturity Date, or upon the closing of the Next Equity Financing, the Company '
    'issues any convertible notes principally for financing purposes other than the Notes (the \u201cAdditional '
    'Notes\u201d), the Note holder shall have the option to purchase Additional Notes the principal of which equals '
    'the Note holder\u2019s pro rata portion of the aggregate principal of Additional Notes sold, based on the '
    'aggregate principal of the Notes held by the Note holder.'
)

subsection_header("6.4", "Successors and Assigns.")
body(
    'Except as otherwise provided herein, the terms and conditions of this Agreement shall inure to the benefit of '
    'and be binding upon the respective successors and assigns of the parties; provided, however, that the Company '
    'may not assign its obligations under this Agreement without the written consent of the Majority Note Holders.  '
    'Nothing in this Agreement, express or implied, is intended to confer upon any party other than the parties '
    'hereto or their respective successors and assigns any rights, remedies, obligations or liabilities under or by '
    'reason of this Agreement, except as expressly provided in this Agreement.'
)

subsection_header("6.5", "Governing Law; Submission to Jurisdiction.")
body(
    'This Agreement will be governed by and construed in accordance with the laws of the State of New York, without '
    'regard to its conflicts of law principles.  Each party irrevocably and unconditionally submits to the exclusive '
    'jurisdiction of the courts of the State of New York and of the federal courts sitting in the City of New York, '
    'County of New York, State of New York, in all actions or proceedings arising out of or relating to this '
    'Agreement, and agrees that all such actions or proceedings must be litigated exclusively in any such court.  '
    'Each party irrevocably waives any objection which it may now or hereafter have to the laying of the venue of '
    'any such litigation in any such court.  Each party further irrevocably waives any right to a trial by jury in '
    'any action or proceeding directly or indirectly arising out of or relating to this Agreement.'
)

subsection_header("6.6", "Counterparts.")
body(
    'This Agreement may be executed in any number of counterparts, each of which when so executed and delivered will '
    'be deemed an original, and all of which together shall constitute one and the same agreement.'
)

subsection_header("6.7", "Facsimile Signatures.")
body(
    'This Agreement may be executed and delivered by facsimile or PDF and upon such delivery each facsimile or PDF '
    'signature will be deemed to have the same effect as if the original signature had been delivered to each other '
    'party.  The original signature copy shall be delivered to each other party by express overnight delivery.  The '
    'failure to deliver the original signature copy and/or the nonreceipt of the original signature copy shall have '
    'no effect upon the binding and enforceable nature of this Agreement.'
)

subsection_header("6.8", "Titles and Headings.")
body(
    'The titles, captions and headings of this Agreement and the Notes are included for ease of reference only and '
    'will be disregarded in interpreting or construing this Agreement and the Notes.  Unless otherwise specifically '
    'stated, all references herein to \u201csections\u201d and \u201cschedules\u201d will mean \u201csections\u201d '
    'and \u201cschedules\u201d to this Agreement.'
)

subsection_header("6.9", "Notices.")
body(
    'Any and all notices required or permitted to be given to a party pursuant to the provisions of this Agreement '
    'or any Note will be in writing and will be effective and deemed to provide such party sufficient notice under '
    'this Agreement or such Note on the earliest of the following:  (a) at the time of personal delivery, if '
    'delivery is in person; (b) one business day after deposit with an express overnight courier for United States '
    'deliveries, or two business days after such deposit for deliveries outside of the United States; (c) three '
    'business days after deposit in the United States mail by certified mail (return receipt requested) for United '
    'States deliveries; or (d) by electronic mail at the time of confirmation back of delivery.  All notices not '
    'delivered personally will be sent with postage and/or other charges prepaid and properly addressed to the '
    'party to be notified.  All communications shall be sent to the respective parties at the following addresses '
    '(or at such other addresses as shall be specified by notice given in accordance with this Section\u00a06.9):'
)
body('If to the Company:')
body('[COMPANY NAME], INC.\n[Address]\n[City, State, Zip]\nEmail:  [_____]', indent=True)
body('If to the Lender:')
body('At the address shown on the Signature Page hereto.', indent=True)

subsection_header("6.10", "Finder\u2019s Fee.")
body(
    'Each party represents that it neither is nor will be obligated for any finder\u2019s fee or commission in '
    'connection with the transactions contemplated by this Agreement.  The Lender agrees to indemnify and to hold '
    'harmless the Company from any liability for any commission or compensation in the nature of a finder\u2019s '
    'fee (and the costs and expenses of defending against such liability or asserted liability) for which such '
    'Lender or any of its officers, partners, employees or representatives is responsible.  The Company agrees to '
    'indemnify and hold harmless the Lender from any liability for any commission or compensation in the nature of '
    'a finder\u2019s fee (and the costs and expenses of defending against such liability or asserted liability) for '
    'which the Company or any of its officers, employees or representatives is responsible.'
)

subsection_header("6.11", "Entire Agreement; Amendments and Waivers.")
body(
    'This Agreement, the Notes and the other documents delivered pursuant hereto constitute the full and entire '
    'understanding and agreement between the parties to this Agreement with regard to the subjects hereof and '
    'thereof.  The Company\u2019s agreements with each of the Lenders are separate agreements, and the sales of '
    'the Notes to each of the Lenders are separate sales.  Nonetheless, any term of this Agreement or the Notes '
    'may be amended and the observance of any term of this Agreement or the Notes may be waived (either generally '
    'or in a particular instance and either retroactively or prospectively), with the written consent of the '
    'Company and the Majority Note Holders.  Any waiver or amendment effected in accordance with this '
    'Section\u00a06.11 shall be binding upon each party to the Note Purchase Agreements and any holder of any '
    'Notes purchased under the Note Purchase Agreements at the time outstanding and each future holder of all '
    'such Notes.'
)

subsection_header("6.12", "Effect of Amendment or Waiver.")
body(
    'The Lender acknowledges that by the operation of Section\u00a06.11 hereof, the Majority Note Holders will '
    'have the right and power to diminish or eliminate all rights of such Lender under this Agreement and each '
    'Note issued to such Lender.'
)

subsection_header("6.13", "Stock Purchase Agreement.")
body(
    'Each Lender understands and agrees that the conversion of the Notes into Conversion Shares may require such '
    'Lender\u2019s execution of certain agreements in the form agreed to by investors in the Next Equity Financing '
    'or the other parties to any Corporate Transaction, relating to the purchase and sale of such securities, as '
    'well as registration, co-sale, rights of first refusal, rights of first offer and voting rights, if any, '
    'relating to such securities.'
)

subsection_header("6.14", "Further Assurances.")
body(
    'Each Lender agrees to execute such further documents and instruments and to take such further actions as may '
    'be reasonably necessary to carry out the purposes and intent of this Agreement and the Note.'
)

subsection_header("6.15", "Expenses.")
body(
    'Each party shall pay all costs and expenses that it incurs with respect to the negotiation, execution, '
    'delivery and performance of this Agreement.'
)

# ============================================================
# SIGNATURE PAGE
# ============================================================
doc.add_page_break()

p_sig = doc.add_paragraph()
p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sig = p_sig.add_run('[Signature Page Follows]')
set_font(r_sig, size=11)

doc.add_paragraph()

p_wit = doc.add_paragraph()
p_wit.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r_wit = p_wit.add_run(
    'IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.'
)
set_font(r_wit, size=11)

doc.add_paragraph()

# Company block
p_c = doc.add_paragraph()
r_c = p_c.add_run('[COMPANY NAME], INC.')
set_font(r_c, bold=True)

for label in ['By:', 'Name:', 'Title:', 'Address:', 'Email:']:
    p_l = doc.add_paragraph()
    r_l = p_l.add_run(f'{label}  ____________________________')
    set_font(r_l)
    p_l.paragraph_format.space_before = Pt(2)
    p_l.paragraph_format.space_after = Pt(2)

doc.add_paragraph()

# Lender block
p_ln = doc.add_paragraph()
r_ln = p_ln.add_run('LENDER:')
set_font(r_ln, bold=True)

for label in ['By:', 'Print Name:', 'Title:', 'Address:', 'Email:', 'Principal Amount of Notes Purchased:']:
    p_l = doc.add_paragraph()
    r_l = p_l.add_run(f'{label}  ____________________________')
    set_font(r_l)
    p_l.paragraph_format.space_before = Pt(2)
    p_l.paragraph_format.space_after = Pt(2)

# ============================================================
# SCHEDULE A
# ============================================================
doc.add_page_break()

heading("SCHEDULE A", level=1, center=True)
heading("FORM OF NOTE", level=1, center=True)

body('(See Attached Convertible Promissory Note)')

doc.add_paragraph()
body(
    'Note:  The Convertible Promissory Note attached as Schedule\u00a0A shall be governed by the laws of the '
    'State of New York and shall include, at minimum, the following terms:'
)
note_terms = [
    'Principal Amount: $[__________]',
    'Issue Date: __________',
    'Maturity Date: __________ (18\u201324 months from Issue Date)',
    'Interest Rate: [__]% per annum, simple interest, not compounding (not to exceed 16% per annum in accordance with New York civil usury limits)',
    'Conversion: as set forth in Section\u00a02.2 of the Agreement',
    'Discount Rate: [__]%',
    'Valuation Cap: $[__________]',
    'No Prepayment without consent of the holder',
    'Governing Law: State of New York',
]
for term in note_terms:
    p_nt = doc.add_paragraph(style='List Bullet')
    r_nt = p_nt.add_run(term)
    set_font(r_nt)

out_path = r'C:\Users\ezaza\OneDrive\Desktop\scopic\ai-slop-fixer\NPA_Output_NY.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
