"""
agents.py
Multi-agent pipeline for fixing AI-slop option agreements.

Pipeline
--------
1. TextExtractor  – pulls plain text from a .docx file (no LLM)
2. ExtractorAgent – LLM call that reads the slop document and returns a
                    structured JSON of all deal-specific terms.
3. WriterAgent    – LLM call that reads the template paragraph/table list
                    + the extracted deal terms, then returns a replacement
                    map: {paragraph_id → new text} and {table_id → new rows}.
"""

from __future__ import annotations

import json
import os
from docx import Document as DocxDocument
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


# ---------------------------------------------------------------------------
# Utility
# ---------------------------------------------------------------------------

def extract_text_from_docx(path: str) -> str:
    """Return the full plain text of a .docx file (paragraphs + tables)."""
    doc = DocxDocument(path)
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Agent 1 – ExtractorAgent
# ---------------------------------------------------------------------------

_EXTRACTOR_SYSTEM = """You are a senior legal analyst specialising in Canadian corporate securities law.

Your task: carefully read the provided stock-option agreement (which may be poorly structured or contain "AI slop") and extract EVERY deal-specific detail into a single, comprehensive JSON object.

Required fields (use null where information is absent):
- company_name             : full legal name of the granting company
- company_jurisdiction     : province/country of incorporation
- optionee_name            : name of the option recipient (or placeholder text)
- optionee_role            : "employee" | "consultant" | "advisor" | other
- grant_date               : date of the grant (or placeholder)
- base_agreement_name      : name of the underlying agreement (e.g. "Employment Agreement", "Consulting Agreement")
- base_agreement_date      : date of the underlying agreement
- option_type              : e.g. "nonqualified stock option", "incentive stock option"
- share_class              : e.g. "Class A Common Shares", "common shares"
- total_shares             : number of shares or formula
- option_price_description : how the exercise price is set
- option_term_years        : term in years (integer)
- vesting_schedule         : full description of how and when options vest
- vesting_cliff_months     : cliff period in months (0 if no cliff)
- vesting_frequency        : "monthly" | "quarterly" | "annually" | other
- vesting_total_months     : total vesting period in months
- post_termination_exercise_months : exercise window after termination (integer)
- cause_definition         : verbatim or paraphrased definition of "Cause"
- acceleration_triggers    : list of corporate events that cause acceleration
- governing_province       : governing law jurisdiction
- lock_up_applicable       : true/false
- shareholder_agreement_applicable : true/false
- equity_accrual_formula   : any hourly/accrual formula (null if not applicable)
- special_provisions       : list of any unique or unusual provisions

Be thorough. Do NOT summarise or omit details."""


def run_extractor_agent(slop_text: str) -> dict:
    """
    Agent 1: Extract all deal-specific terms from the slop document.

    Returns a dict with the structured deal terms.
    """
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": _EXTRACTOR_SYSTEM},
            {
                "role": "user",
                "content": (
                    "Extract all deal terms from this option agreement:\n\n"
                    + slop_text
                ),
            },
        ],
        response_format={"type": "json_object"},
        temperature=0,
    )
    return json.loads(response.choices[0].message.content)


# ---------------------------------------------------------------------------
# Agent 2 – WriterAgent  (generic — works for any legal document type)
# ---------------------------------------------------------------------------

_WRITER_SYSTEM = """You are an expert legal document drafter specialising in Canadian corporate and commercial law.

Your task: given (a) the full paragraph-by-paragraph skeleton of a professionally formatted legal template, and (b) deal terms extracted from an AI-generated "slop" document, produce a COMPLETE replacement map that transforms the template into a properly populated, deal-specific legal document.

=== MANDATORY RULES ===

THOROUGHNESS — SCAN EVERY PARAGRAPH
1. Examine EVERY paragraph in the template list without exception.
2. For any paragraph that contains a placeholder, generic party label, blank field, or stale generic term — replace it.
3. Placeholder patterns to detect and replace:
   - Square-bracket tags: [COMPANY], [NAME], [ADDRESS], [DATE], [NUMBER], [TITLE], [x], etc.
   - Angle-star tokens: <*>
   - Blank underscores: _____ or ____________
   - Inline placeholder text inside sentences (e.g. "between ________ (the Corporation)")
4. Do NOT skip any paragraph — missing even one placeholder is a failure.

STRUCTURE & STYLE
5. Preserve the exact paragraph numbering and section order of the template.
6. Keep each paragraph's Word style as-is — only the text changes.
7. Never introduce bullet points where the template uses prose.
8. Use the same formal legal language and sentence construction as the template.
9. Do NOT add new paragraphs or sections beyond what exists in the template.

PARTY NAMES & REFERENCES
10. Replace ALL generic company/party placeholders with actual names from the deal terms.
11. Replace ALL generic person-name placeholders (signatories, addressees) with actual names.
12. Update all addresses, dates, titles, roles, amounts, and percentages with actual values.
13. If the underlying agreement type differs from the template's default (e.g. Consulting vs. Employment), update ALL references consistently throughout every paragraph — no partial updates.

SIGNATURE BLOCKS
14. Fill in all signature blocks (company name, signatory name, title, date) from deal terms.
15. Fill in the optionee/counterparty name in their signature block.

TABLES
16. For every table, fill in cells that correspond to deal-specific data (party names, addresses, dates, share counts, prices, etc.).

SCHEDULE A / NOTICE FORMS (if present)
17. Pre-fill any "Notice of Exercise" or similar schedule with the known fixed values:
    - Grant date, exercise price, total shares subject to option.
    - Leave fields that are completed at exercise time (number to purchase, total payable, delivery address) blank.

POST-PROCESSING AWARENESS
18. After producing the map, mentally re-read each replacement you generated and confirm it contains no leftover placeholders, generic party names, or stale template references.

OUTPUT FORMAT — return exactly:
{
  "paragraphs": { "<id_as_string>": "<full replacement text>", ... },
  "tables": { "<table_id_as_string>": [["r0c0","r0c1",...], ["r1c0",...], ...], ... }
}

Include EVERY paragraph and table that needs a change. Omit only paragraphs that are already correct.
Preserve any leading tabs (\\t) that appear at the start of the original paragraph text.
"""


def run_writer_agent(
    template_paragraphs: list[dict],
    template_tables: list[dict],
    deal_terms: dict,
    slop_text: str,
) -> dict:
    """
    Agent 2: Produce the full replacement map for all paragraphs and tables.

    Returns::
        {
          "paragraphs": {"<id>": "<new text>", ...},
          "tables":     {"<id>": [[row0col0, ...], [row1col0, ...]], ...}
        }
    """
    para_context = json.dumps(
        [
            {"id": p["id"], "style": p["style"], "text": p["text"]}
            for p in template_paragraphs
            if p["has_content"]
        ],
        indent=2,
        ensure_ascii=False,
    )

    user_content = (
        "Return your answer as a JSON object.\n\n"
        "=== TEMPLATE PARAGRAPHS (id, style, current text) ===\n"
        + para_context
        + "\n\n=== TEMPLATE TABLES ===\n"
        + json.dumps(template_tables, indent=2, ensure_ascii=False)
        + "\n\n=== DEAL TERMS (extracted from slop) ===\n"
        + json.dumps(deal_terms, indent=2, ensure_ascii=False)
        + "\n\n=== ORIGINAL SLOP DOCUMENT (for reference) ===\n"
        + slop_text
    )

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": _WRITER_SYSTEM},
            {"role": "user", "content": user_content},
        ],
        response_format={"type": "json_object"},
        temperature=0,
        max_tokens=8192,
    )
    return json.loads(response.choices[0].message.content)
