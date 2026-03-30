import glob
from docx import Document

files = {f.split("\\")[-1]: f for f in glob.glob(r"c:\Users\ezaza\OneDrive\Desktop\scopic\ai-slop-fixer\public\*.docx")}

SLOP = files["Dentons Use Case - Option Agreement AI Slop.docx"]
TMPL = files["Option Agreement template.docx"]
OUT  = files["Option Agreement - ABC Startup Corp.docx"]

print("=== SLOP: ALL CONTENT ===")
doc = Document(SLOP)
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"  [{i}]: {p.text}")

print()
print("=== TEMPLATE: placeholder / blank fields only ===")
doc2 = Document(TMPL)
for i, p in enumerate(doc2.paragraphs):
    if ("_" in p.text or "[" in p.text) and p.text.strip():
        print(f"  [{i}]: {p.text[:120]}")
for i, tbl in enumerate(doc2.tables):
    for r, row in enumerate(tbl.rows):
        for c, cell in enumerate(row.cells):
            if ("_" in cell.text or "[" in cell.text) and cell.text.strip():
                print(f"  Table{i}[{r},{c}]: {cell.text[:120]}")

print()
print("=== OUTPUT: remaining blanks ===")
doc3 = Document(OUT)
for i, p in enumerate(doc3.paragraphs):
    if ("_" in p.text or "[" in p.text) and p.text.strip():
        print(f"  [{i}]: {p.text[:120]}")
