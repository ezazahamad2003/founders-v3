from __future__ import annotations

import os
import re
from dataclasses import dataclass
from pathlib import Path

from dotenv import load_dotenv
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from openai import OpenAI

# Load .env file (repo-root or current working dir)
load_dotenv()

MODEL = "gpt-5.1"
MAX_TOKENS = 4096

_client: OpenAI | None = None


def get_client() -> OpenAI:
    global _client
    if _client is None:
        api_key = os.getenv("OPENAI_API_KEY")
        _client = OpenAI(api_key=api_key)
    return _client


BASE_DIR = Path(__file__).resolve().parent
PUBLIC_DIR = BASE_DIR / "public"
OUTPUTS_DIR = BASE_DIR / "outputs"

MAX_CHARS_PER_FILE = 50000  # keep prompts bounded
EXCERPT_CHARS_FOR_LOGGING = 800  # included in output files for debugging only


def load_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def list_agreement_files() -> list[Path]:
    if not PUBLIC_DIR.exists():
        raise FileNotFoundError(f"Missing directory: {PUBLIC_DIR}")
    files = sorted([p for p in PUBLIC_DIR.iterdir() if p.is_file() and p.suffix.lower() == ".docx"])
    if not files:
        raise FileNotFoundError(f"No .docx files found in: {PUBLIC_DIR}")
    return files


def extract_docx_text(path: Path) -> str:
    try:
        document = DocxDocument(path)  # python-docx can read from a file path
    except (PackageNotFoundError, ValueError, KeyError) as exc:
        return f""

    parts: list[str] = []
    for paragraph in document.paragraphs:
        content = (paragraph.text or "").strip()
        if content:
            parts.append(content)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    return text


def normalize_ws(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def truncate(text: str, max_chars: int) -> str:
    if not text:
        return ""
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "…"


@dataclass(frozen=True)
class Question:
    qid: str
    title: str
    body: str

    def as_prompt_text(self) -> str:
        if self.title:
            return f"{self.qid} — {self.title}\n\n{self.body}".strip()
        return f"{self.qid}\n\n{self.body}".strip()


def parse_questions(text: str) -> list[Question]:
    """
    Expected format (blank lines allowed):
      Q1 — Title

      Body...

      Q2 — Title
      Body...
    """
    lines = [l.rstrip() for l in (text or "").splitlines()]
    items: list[Question] = []

    current_qid: str | None = None
    current_title: str = ""
    current_body_lines: list[str] = []

    header_re = re.compile(r"^(Q\d+)\s*(?:—\s*(.*))?$")

    def flush():
        nonlocal current_qid, current_title, current_body_lines
        if not current_qid:
            return
        body = "\n".join([l for l in current_body_lines]).strip()
        items.append(Question(qid=current_qid, title=(current_title or "").strip(), body=body))
        current_qid = None
        current_title = ""
        current_body_lines = []

    for raw in lines:
        line = raw.strip()
        if not line:
            if current_qid:
                current_body_lines.append("")
            continue
        m = header_re.match(line)
        if m:
            flush()
            current_qid = m.group(1)
            current_title = (m.group(2) or "").strip()
            continue
        if current_qid:
            current_body_lines.append(line)

    flush()

    # Back-compat: if file still uses "1." style numbering, parse into Q1..Qn
    if not items:
        numbered: list[str] = []
        current: list[str] = []
        for line in [l.strip() for l in (text or "").splitlines() if l.strip()]:
            if line[0].isdigit() and "." in line:
                if current:
                    numbered.append(" ".join(current))
                    current = []
            current.append(line)
        if current:
            numbered.append(" ".join(current))
        items = [
            Question(qid=f"Q{i}", title="", body=q) for i, q in enumerate(numbered, start=1)
        ]

    if not items:
        raise ValueError("No questions found in questions.txt")

    return items


def ask_model(system_prompt: str | None, user_prompt: str) -> str:
    messages = []

    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})

    messages.append({"role": "user", "content": user_prompt})

    response = get_client().chat.completions.create(
        model=MODEL,
        messages=messages,
        max_completion_tokens=MAX_TOKENS,
    )

    content = response.choices[0].message.content
    return content.strip() if content and content.strip() else "[EMPTY MODEL OUTPUT]"


def build_user_prompt(*, agreement_filename: str, agreement_excerpt: str, question: Question) -> str:
    excerpt_note = (
        f"(Excerpt length: {len(agreement_excerpt)} chars; may be truncated.)"
        if agreement_excerpt
        else "(No extractable text found.)"
    )
    return (
        "You are reviewing the Agreement below. Use ONLY the provided Agreement text; "
        "if something is missing/unclear, say so.\n\n"
        f"Agreement file: {agreement_filename}\n"
        f"{excerpt_note}\n\n"
        "=== AGREEMENT TEXT (EXCERPT) ===\n"
        f"{agreement_excerpt}\n"
        "=== END AGREEMENT TEXT ===\n\n"
        "TASK:\n"
        f"{question.as_prompt_text()}\n"
    )


def run_eval_to_file(
    *,
    title: str,
    out_path: Path,
    system_prompt: str | None,
    questions: list[Question],
    agreement_files: list[Path],
) -> None:
    total = len(agreement_files) * len(questions)
    idx = 0

    with out_path.open("w", encoding="utf-8") as f:
        f.write(title + "\n")
        f.write("=" * len(title) + "\n\n")

        for file_path in agreement_files:
            raw_text = extract_docx_text(file_path)
            normalized = normalize_ws(raw_text)
            agreement_excerpt = truncate(normalized, MAX_CHARS_PER_FILE)
            logging_excerpt = truncate(normalized, EXCERPT_CHARS_FOR_LOGGING)

            for q in questions:
                idx += 1
                print(f"[{idx}/{total}] {file_path.name} :: {q.qid} ...")

                user_prompt = build_user_prompt(
                    agreement_filename=file_path.name,
                    agreement_excerpt=agreement_excerpt,
                    question=q,
                )
                answer = ask_model(system_prompt, user_prompt)

                f.write(f"Item {idx}/{total}\n")
                f.write(f"File: {file_path.name}\n")
                f.write(f"Question: {q.qid} — {q.title}\n\n")
                f.write("Prompt:\n")
                f.write(q.body + "\n\n")
                f.write(f"Agreement excerpt used (first {EXCERPT_CHARS_FOR_LOGGING} chars):\n")
                f.write(logging_excerpt + "\n\n")
                f.write("Answer:\n")
                f.write(answer + "\n\n")
                f.write("-" * 60 + "\n\n")
                f.flush()  # keep partial outputs if interrupted


def main():
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError(
            "OPENAI_API_KEY is not set. Set it in your environment or in a .env file."
        )

    questions_text = load_text_file(BASE_DIR / "questions.txt")
    system_prompt = load_text_file(BASE_DIR / "system_prompt.txt")

    questions = parse_questions(questions_text)
    agreement_files = list_agreement_files()
    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)

    # WITH system prompt
    with_prompt_out = OUTPUTS_DIR / "with_prompt.txt"
    run_eval_to_file(
        title="EVAL RUN - WITH SYSTEM PROMPT (Q1–Q5 × 5 DOCX)",
        out_path=with_prompt_out,
        system_prompt=system_prompt,
        questions=questions,
        agreement_files=agreement_files,
    )

    # WITHOUT system prompt
    without_prompt_out = OUTPUTS_DIR / "without_prompt.txt"
    run_eval_to_file(
        title="EVAL RUN - WITHOUT SYSTEM PROMPT (Q1–Q5 × 5 DOCX)",
        out_path=without_prompt_out,
        system_prompt=None,
        questions=questions,
        agreement_files=agreement_files,
    )

    print("Eval complete.")
    print("Generated:")
    print(f"- {with_prompt_out}")
    print(f"- {without_prompt_out}")


if __name__ == "__main__":
    main()
